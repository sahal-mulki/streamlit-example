from openai import OpenAI
import os

from docx import Document
from docx.shared import Inches
from docx.shared import Pt

def create_worksheet_layout(topic: str, audience: str, n_questions: int, openaikey: str):
  client = OpenAI(api_key=openaikey)

  prompt = """Make a layout for a question worksheet. Be sure to make it engaging and use interesting, long questions. ADD LINES OF SPACE FOR HOW MANY LINES THE QUESTION WOULD NEED FOR ANSWERING IN A PDF. FILL IN THE ANSWER FIELD WITH THE CORRECT ANSWER FOR THE QUESTION.
              It should be in this format:
              {"questions": {"number": 1, "question": "What's a dog?", "answer": "A dog is an animal which......", "num-lines-for-space": "Number"}}
              ADHERE TO THE JSON FORMAT AND DO NOT PROVIDE ANY ADDITIONAL INFORMATION. Do not use any controversial or political statements in the presentation. DO NOT USE EMOJIS OR ANY SPECIAL CHARACTERS. CREATE AS MANY QUESTIONS AS SPECIFIED."""

  response = client.chat.completions.create(
    model="gpt-3.5-turbo-1106",
    response_format={ "type": "json_object" },
    messages=[
      {
        "role": "system",
        "content": prompt
      },
      {
        "role": "user",
        "content": "Topic: " + topic + " , Audience: " + audience + ",  Create" + str(n_questions) + "Questions" 
      }
    ],
    temperature=0.3,
    frequency_penalty=0.22,
    max_tokens=2000,
    top_p=1
  )

  exec("questions = " + response.choices[0].message.content, globals())

  return questions

def process_output_worksheet(output: dict, n_questions: int):
  questions = []
  questions_spaces = []
  answers = []


  for x in range(n_questions):
    questions.append(output["questions"][x]["question"])

  for x in range(n_questions):
    questions_spaces.append(output["questions"][x]["num-lines-for-space"])

  for x in range(n_questions):
    answers.append(output["questions"][x]["answer"])

  return questions, questions_spaces, answers

def make_sheet_title(topic: str, audience: str, openaikey: str):
  client = OpenAI(api_key=openaikey)

  response = client.chat.completions.create(
    model="gpt-3.5-turbo-1106",
    messages=[
      {
        "role": "system",
        "content": "Make a heading for a question worksheet. DO NOT USE EMOJIS OR SPECIAL CHARACTERS."
      },
      {
        "role": "user",
        "content": "Topic: " + topic + " , Audience: " + audience
      }
    ],
    temperature=0.4,
    frequency_penalty=0.22,
    max_tokens=100,
    top_p=1
  )

  return response.choices[0].message.content

def make_question_sheet(title: str, questions: list, questions_spaces: list):
  document = Document()

  title1 = document.add_heading(title, 0)
  title1.font = 'Helvetica'

  document.add_paragraph("")


  font = document.styles['Normal'].font
  font.name = 'Helvetica'
  font.size = Pt(13)

  poo = -1

  for x in questions:
    poo += 1
    document.add_paragraph(str(poo + 1) + ". " + x)

    for y in range(int(questions_spaces[poo])):
      document.add_paragraph("___________________________________________________________")

  document.save(os.path.join('generated-question-sheets/', str(title + " - QUESTION SHEET" + '.docx')))

  return os.path.join('generated-question-sheets/', str(title + " - QUESTION SHEET" + '.docx'))

def make_answer_sheet(title: str, questions: list, answers: list):
  document = Document()

  title1 = document.add_heading(title, 0)
  title1.font = 'Helvetica'

  document.add_paragraph("")


  font = document.styles['Normal'].font
  font.name = 'Helvetica'
  font.size = Pt(13)

  poo = -1

  for x in questions:
    poo += 1
    document.add_paragraph(str(poo + 1) + ". " + x)

    document.add_paragraph("Answer: " + answers[poo])

  document.save(os.path.join('generated-answer-sheets/', str(title + " - ANSWER SHEET" + '.docx')))

  return os.path.join('generated-answer-sheets/', str(title + " - ANSWER SHEET" + '.docx'))
